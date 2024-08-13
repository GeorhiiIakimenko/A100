from typing import Optional, Dict, Any, List, Tuple
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import Chroma
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
import sqlite3
import time
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.prompts.chat import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import AIMessage, HumanMessage, SystemMessage
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, Form, Request, Body
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, JSONResponse
from starlette.staticfiles import StaticFiles
from passlib.context import CryptContext
from starlette.middleware.sessions import SessionMiddleware
import uvicorn
import boto3
from fnmatch import fnmatchcase
import json
from pydantic import BaseModel
import os
from docx.oxml.ns import qn
from docx import Document as DocxDocument
from io import BytesIO
from authlib.integrations.starlette_client import OAuth
from dotenv import load_dotenv
import random
import string
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

GOOGLE_CLIENT_ID = os.getenv("add")
GOOGLE_CLIENT_SECRET = os.getenv("add")
SECRET_KEY = os.getenv("123456789")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not OPENAI_API_KEY:
    raise ValueError("Missing OPENAI_API_KEY environment variable")

os.environ['OPENAI_API_KEY'] = OPENAI_API_KEY

model_name = "gpt-4o"
temperature = 0
llm = ChatOpenAI(model=model_name, temperature=temperature)
embeddings = OpenAIEmbeddings()

current_user = 'A100'

# Настройка клиента для Yandex S3
session = boto3.session.Session()
s3_client = session.client(
    service_name='s3',
    endpoint_url='https://storage.yandexcloud.net',
    aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
    aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY')
)

CHROMA_PATH = f'./chroma/{current_user}/'

oauth = OAuth()
oauth.register(
    name='google',
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    authorize_url='https://accounts.google.com/o/oauth2/auth',
    authorize_params=None,
    access_token_url='https://accounts.google.com/o/oauth2/token',
    access_token_params=None,
    refresh_token_url=None,
    redirect_uri='http://localhost:8222/login',
    client_kwargs={'scope': 'openid profile email'},
)

def init_metadata_db():
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()

        # Create uploaded_docs table if not exists
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS uploaded_docs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            global_source TEXT,
            filename TEXT
        );
        ''')

        # Create history_messages table if not exists
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS history_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id TEXT,
            user_type TEXT,
            message TEXT,
            chat_id TEXT,
            tmstmp DATETIME DEFAULT CURRENT_TIMESTAMP
        );
        ''')

        # Create chats table if not exists
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS chats (
            chat_id TEXT PRIMARY KEY,
            chat_name TEXT
        );
        ''')

        # Check and add columns to history_messages table if needed
        cursor.execute("PRAGMA table_info(history_messages);")
        columns = [column[1] for column in cursor.fetchall()]
        if 'chat_id' not in columns:
            try:
                cursor.execute('''
                ALTER TABLE history_messages ADD COLUMN chat_id TEXT;
                ''')
                print("chat_id column added successfully")
            except sqlite3.OperationalError:
                print("chat_id column already exists or cannot be added")

        # Check if users table exists and has the email and google_id columns
        cursor.execute("PRAGMA table_info(users);")
        columns = [column[1] for column in cursor.fetchall()]

        if 'email' not in columns or 'google_id' not in columns:
            # Rename old users table
            cursor.execute("ALTER TABLE users RENAME TO users_old;")

            # Create new users table with email and google_id columns
            cursor.execute('''
                CREATE TABLE users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT UNIQUE,
                    hashed_password TEXT,
                    email TEXT UNIQUE,
                    google_id TEXT UNIQUE
                );
            ''')

            # Copy data from old users table to new users table
            cursor.execute('''
                INSERT INTO users (id, username, hashed_password)
                SELECT id, username, hashed_password FROM users_old;
            ''')

            cursor.execute('''
            CREATE TABLE IF NOT EXISTS chat_names (
                chat_id TEXT PRIMARY KEY,
                chat_name TEXT
            );
            ''')

            # Drop old users table
            cursor.execute("DROP TABLE users_old;")
            print("email and google_id columns added successfully by recreating the table")
        else:
            print("email and google_id columns already exist")

        conn.commit()

init_metadata_db()

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def get_password_hash(password):
    return pwd_context.hash(password)

def add_user_to_db(username: str, password: str = None, email: str = None, google_id: str = None):
    hashed_password = get_password_hash(password) if password else None
    try:
        with sqlite3.connect('metadata.db') as conn:
            conn.execute("INSERT INTO users (username, hashed_password, email, google_id) VALUES (?, ?, ?, ?)",
                         (username, hashed_password, email, google_id))
            conn.commit()
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="User already registered")

def authenticate_user(username: str, password: str):
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT hashed_password FROM users WHERE username = ?", (username,))
        row = cursor.fetchone()
    if row and pwd_context.verify(password, row[0]):
        return True
    return False

class SQLiteChatHistory():
    def __init__(self, db_path="metadata.db"):
        self.db_path = db_path

    def add_message(self, message, chat_id, chat_name=None):
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        if isinstance(message, HumanMessage):
            user_type = "human"
            message = message.content
        elif isinstance(message, AIMessage):
            user_type = "ai"
            message = message.content
        elif isinstance(message, SystemMessage):
            user_type = "system"
            message = message.content
        else:
            raise ValueError("Invalid message type")
        c.execute("INSERT INTO history_messages (user_id, user_type, message, chat_id) VALUES (?, ?, ?, ?)",
                  (current_user, user_type, message, chat_id))
        if chat_name:
            c.execute("INSERT OR IGNORE INTO chat_names (chat_id, chat_name) VALUES (?, ?)", (chat_id, chat_name))
            c.execute("UPDATE chat_names SET chat_name = ? WHERE chat_id = ?", (chat_name, chat_id))
        conn.commit()
        conn.close()

    def messages(self, chat_id):
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(f"SELECT * FROM history_messages WHERE chat_id = ? ORDER BY id", (chat_id,))
        resp = c.fetchall()
        chat_history = []
        for row in resp:
            id, user_id, user_type, message, chat_id, tmstmp = row
            if user_type == "human":
                chat_history.append(HumanMessage(content=message))
            elif user_type == "ai":
                chat_history.append(AIMessage(content=message))
            elif user_type == "system":
                chat_history.append(SystemMessage(content=message))
        conn.close()
        return ChatMessageHistory(messages=chat_history)

    def delete_chat_history_last_n(self, chat_id, n=10):
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(f'''
        WITH max_id AS (SELECT MAX(id) AS maxid FROM history_messages WHERE user_id = '{current_user}' AND chat_id = '{chat_id}')
        DELETE FROM history_messages
        WHERE id BETWEEN (SELECT maxid FROM max_id) - {n} AND (SELECT maxid FROM max_id)
        ''')
        conn.commit()
        conn.close()



def add_filename_to_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''INSERT INTO uploaded_docs (global_source, filename) VALUES ('{source}', '{filename}') ; ''')

def delete_filename_from_metadata(source, filename):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute(f'''DELETE FROM uploaded_docs WHERE global_source = '{source}' AND filename ='{filename}' ; ''')

class Document:
    def __init__(self, source: str, page_content: str, metadata: Optional[Dict[str, Any]] = None):
        self.source = source
        self.page_content = page_content
        self.metadata = metadata if metadata is not None else {'source': source}

def get_uploaded_filenames(source) -> List[str]:
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute(f"SELECT filename FROM uploaded_docs WHERE global_source = ?", (source,))
        rows = cursor.fetchall()
    filenames = [row[0] for row in rows]
    return filenames

def load_s3_files(bucket: str, prefix: str, suffix: str) -> List[str]:
    """List files in a given S3 bucket with a specified prefix and suffix."""
    try:
        response = s3_client.list_objects_v2(Bucket=bucket, Prefix=prefix)
        files = [content['Key'] for content in response.get('Contents', []) if content['Key'].endswith(suffix)]
        if not files:
            print(f"No files found in bucket {bucket} with prefix {prefix} and suffix {suffix}")
        else:
            print(f"Files found in bucket {bucket} with prefix {prefix} and suffix {suffix}: {files}")
        return files
    except Exception as e:
        print(f"Error listing files in bucket {bucket} with prefix {prefix} and suffix {suffix}: {e}")
        return []

def load_docx_new(source, bucket: str) -> List[Document]:
    prefix = 'A100/docx/'
    suffix = '.docx'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read()

                # Используем BytesIO для чтения содержимого файла как бинарного потока
                doc_stream = BytesIO(content)
                doc = DocxDocument(doc_stream)

                # Извлекаем текст из документа docx
                full_text = []
                image_counter = 1

                # Получаем имя файла без расширения и создаем соответствующую папку
                filename_without_extension = os.path.splitext(os.path.basename(file))[0]
                image_folder = filename_without_extension  # Используем оригинальное имя файла для папки

                for para in doc.paragraphs:
                    # Обработка параграфов для создания ссылок на изображения
                    para_text = para.text
                    for run in para.runs:
                        for drawing in run.element.findall('.//a:blip', namespaces={
                            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                            image_rId = drawing.get(qn('r:embed'))
                            image_part = doc.part.related_parts[image_rId]
                            image_filename = f'image_{image_counter:02d}.{image_part.content_type.split("/")[-1]}'
                            image_counter += 1

                            # Загрузка изображения в бакет Яндекса
                            img_content = image_part.blob
                            s3_image_key = f"A100/images/{image_folder}/{image_filename}"
                            s3_client.put_object(
                                Bucket=bucket,
                                Key=s3_image_key,
                                Body=img_content,
                                ContentDisposition='inline',
                                ContentType=image_part.content_type
                            )

                            # Генерация URL для изображения
                            s3_image_url = f"https://storage.yandexcloud.net/{bucket}/{s3_image_key}"
                            para_text += f'\n{s3_image_url}'
                    full_text.append(para_text)
                content = '\n'.join(full_text)

                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading docx file {file}: {e}")

    return docs if docs else None

def load_txts(source, bucket: str) -> List[Document]:
    prefix = f'{current_user}/txt/'
    suffix = '.txt'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    docs = []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = obj['Body'].read().decode('utf-8')
                docs.append(Document(source=file, page_content=content))
            except Exception as e:
                print(f"Error reading txt file {file}: {e}")

    return docs if docs else None

def load_jsons(source, bucket: str) -> Tuple[List[Document], List[dict]]:
    prefix = f'{current_user}/json/'
    suffix = '.json'
    files = load_s3_files(bucket, prefix, suffix)
    uniq_files = get_uploaded_filenames(source) or []

    json_docs, json_metadata = [], []
    for file in files:
        if not any(fnmatchcase(file, f"*{pattern}*") for pattern in uniq_files):
            try:
                obj = s3_client.get_object(Bucket=bucket, Key=file)
                content = json.loads(obj['Body'].read().decode('utf-8'))
                json_docs.append(content)
                json_metadata.append({'source': file})
            except Exception as e:
                print(f"Error reading json file {file}: {e}")

    return (json_docs, json_metadata) if json_docs else (None, None)

def load_documents(global_source, bucket: str, file_types: List[str]) -> dict:
    """
    Загружаем документы в зависимости от типа документа из Yandex S3
    """
    all_docs = {'txt': None, 'json': None, 'json_metadata': None, 'docx': None}
    if 'txt' in file_types:
        txt_docs = load_txts(global_source, bucket)
        all_docs['txt'] = txt_docs
    if 'json' in file_types:
        json_docs, json_metadata = load_jsons(global_source, bucket)
        all_docs['json'] = json_docs
        all_docs['json_metadata'] = json_metadata
    if 'docx' in file_types:
        docx_docs = load_docx_new(global_source, bucket)
        all_docs['docx'] = docx_docs
    return all_docs

# Пример использования
DATA_BUCKET = 'utlik'
DOCS = load_documents('s3', DATA_BUCKET, ['txt', 'json', 'docx'])

def split_docs_to_chunks(documents: dict, file_types: List[str], chunk_size=2000, chunk_overlap=500):
    all_chunks = []
    if 'txt' in file_types and documents['txt'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['txt']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    if 'json' in file_types and documents['json'] is not None:
        json_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        json_chunks = json_splitter.create_documents([json.dumps(doc, ensure_ascii=False) for doc in documents['json']],
                                                     metadatas=documents['json_metadata'])
        all_chunks.extend(json_chunks)

    if 'docx' in file_types and documents['docx'] is not None:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        txt_chunks = [text_splitter.split_documents([doc]) for doc in documents['docx']]
        txt_chunks = [item for sublist in txt_chunks for item in sublist]
        all_chunks.extend(txt_chunks)

    return all_chunks

chunks_res = split_docs_to_chunks(DOCS, ['txt', 'json', 'docx'])

def get_chroma_vectorstore(documents, embeddings, persist_directory):
    if os.path.isdir(persist_directory) and os.listdir(persist_directory):
        print("Loading existing Chroma vectorstore...")
        vectorstore = Chroma(
            embedding_function=embeddings, persist_directory=persist_directory
        )

        existing_files = get_uploaded_filenames('local')
        uniq_sources_to_add = set(
            doc.metadata['source'] for doc in chunks_res
            if doc.metadata['source'] not in existing_files
        )

        if uniq_sources_to_add:
            vectorstore.add_documents(
                documents=[doc for doc in chunks_res if doc.metadata['source'] in uniq_sources_to_add],
                embedding=embeddings
            )
            for filename in uniq_sources_to_add:
                add_filename_to_metadata('local', filename)
        else:
            print('Новых документов не было, пропускаем шаг добавления')

    else:
        print("Creating and indexing new Chroma vectorstore...")
        vectorstore = Chroma.from_documents(
            documents=documents,
            embedding=embeddings, persist_directory=persist_directory
        )
        uniq_sources_to_add = set(doc.metadata['source'] for doc in documents)
        for filename in uniq_sources_to_add:
            add_filename_to_metadata('local', filename)

    return vectorstore

vectorstore = get_chroma_vectorstore(documents=chunks_res, embeddings=embeddings, persist_directory=CHROMA_PATH)
retriever = vectorstore.as_retriever(search_kwargs={"k": 2}, search_type='similarity')

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

chat_history_for_chain = SQLiteChatHistory()

prompt_new = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            '''
            Контекст:
            Context: 
                Вы - ассистент для ответов на вопросы, специализирующийся на помощи новым сотрудникам с вопросами, связанными с их работой. Ваша задача - предоставлять точную информацию, рекомендации по правильным действиям и использованию программного обеспечения, основываясь на следующих источниках:

Предоставленный контекст {context}
История чата (chat_history), если это необходимо для лучшего понимания проблемы
Содержимое следующих файлов, находящихся в бакете utlik с префиксом A100/docx/ и суффиксом .docx:

01. Приложение 1 Требования к оформлению документов (1).docx


Приложение 2 Шаблоны имени (нейминга) карточек и файлов документов в СЭД.docx




Приложение 3 Запуск и прерывание процессов обработки документов_ред..docx




Приложение 4 Перенос срока выполнения задачи в СЭД (1).docx




Приложение 5 Примерный перечень документов, не подлежащих регистрации.docx




Приложение 10 Входящий договор с ЭЦП, поступивший через Quidiox (1).docx




Приложение 13 Создание контрагента (1).docx




Приложение 14 Направление задачи на изменение кадровой инф-ции в СЭД (1).docx




Приложение 15 Инструкция о получении, продлении ЭЦП.docx


Положение о документообороте v3.0.docx
Приложение-16-Перечень-владельцев-ЭЦП (2).docx
Приложения а100 * 18*47_48_16_17.docx



Инструкции:

Внимательно прочитайте вопрос пользователя и поймите его суть.
Проанализируйте предоставленный контекст {context} и содержимое всех указанных файлов для поиска релевантной информации.
При работе с файлами:
a. Убедитесь, что у вас есть доступ к содержимому каждого файла.
b. Просматривайте содержимое файлов полностью, не пропуская никакой информации.
c. Обратите особое внимание на файлы, связанные с темой вопроса (например, файлы об ЭЦП при вопросах об электронной подписи).
Если ответ найден в контексте или файлах, предоставьте подробный и точный ответ, обязательно указывая источник информации (название файла).
Используйте информацию из всех релевантных файлов, комбинируя ее при необходимости для полного ответа.
Если в контексте есть ссылки на изображения, обязательно отобразите их в вашем ответе.
Используйте информацию из истории чата только для лучшего понимания контекста вопроса, но не как основной источник ответа.
Не ищите информацию в интернете. Используйте только предоставленные источники.
Если ответ не найден в контексте или файлах после тщательного поиска, ответьте: "Из представленного контекста и доступных документов ответа нет".
Если вопрос касается конкретной процедуры или использования программного обеспечения, предоставьте пошаговые инструкции, основываясь на информации из файлов.
Если информация в разных документах противоречит друг другу, укажите на это и предоставьте информацию из самого актуального документа (если возможно определить).
Если вопрос требует дополнительного уточнения, задайте уточняющие вопросы пользователю.
При ответе на вопросы об ЭЦП обязательно обращайтесь к файлам "15. Приложение 15 Инструкция о получении, продлении ЭЦП.docx" и "Приложение-16-Перечень-владельцев-ЭЦП (2).docx".
            Вопрос:
            Question: {question}
            ''',
        ),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{question}"),
    ]
)

chain_new = prompt_new | llm

chain_with_message_history = RunnableWithMessageHistory(
    chain_new,
    lambda session_id: chat_history_for_chain.messages(session_id),
    input_messages_key="question",
    history_messages_key="chat_history",
)

app = FastAPI()
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)

def generate_random_password(length: int = 8):
    characters = string.ascii_letters + string.digits
    return ''.join(random.choice(characters) for _ in range(length))

def send_reset_password_email(email: str, new_password: str):
    # Здесь должна быть логика для отправки email с новым паролем
    print(f"Отправка нового пароля {new_password} на email {email}")

# Подключение статических файлов
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/register", response_class=HTMLResponse)
async def get_register():
    return FileResponse("static/register.html")

@app.post("/register")
async def post_register(username: str = Form(...), email: str = Form(...), password: str = Form(...)):
    add_user_to_db(username, password, email)
    return RedirectResponse(url="/login", status_code=303)

@app.get("/login", response_class=HTMLResponse)
async def get_login():
    return FileResponse("static/login.html")

@app.post("/login")
async def post_login(username: str = Form(...), password: str = Form(...)):
    if authenticate_user(username, password):
        return RedirectResponse(url="/", status_code=303)
    else:
        return HTMLResponse("Invalid username or password", status_code=400)

@app.get('/auth')
async def auth(request: Request):
    token = await oauth.google.authorize_access_token(request)
    user = await oauth.google.parse_id_token(request, token)
    if user:
        add_user_to_db(user['name'], None, user['email'], user['sub'])
        return RedirectResponse(url="/")
    else:
        return HTMLResponse("Authorization failed", status_code=400)


@app.put("/rename_chat/{chat_id}")
async def rename_chat(chat_id: str, new_title: dict = Body(...)):
    logger.info(f"Attempting to rename chat {chat_id} to {new_title}")
    new_title_str = new_title.get("new_title")
    if not new_title_str:
        logger.error("New title is missing in the request body")
        raise HTTPException(status_code=400, detail="New title is required")

    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("UPDATE chats SET chat_name = ? WHERE chat_id = ?", (new_title_str, chat_id))
        if cursor.rowcount == 0:
            logger.error(f"Chat with id {chat_id} not found in the database")
            raise HTTPException(status_code=404, detail="Chat not found")
        conn.commit()

    logger.info(f"Successfully renamed chat {chat_id} to {new_title_str}")
    return JSONResponse(content={"status": "success", "new_title": new_title_str})

@app.get('/login/google')
async def login_google(request: Request):
    redirect_uri = 'http://localhost:8222/auth'
    return await oauth.google.authorize_redirect(request, redirect_uri)

@app.get("/", response_class=HTMLResponse)
async def read_root():
    return FileResponse("static/index.html")


@app.websocket("/ws/rag_chat/")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    logger.info("WebSocket connection accepted")
    try:
        while True:
            data = await websocket.receive_json()
            logger.info(f"Received data: {data}")
            action = data.get('action')

            if action == "get_history":
                chat_id = data.get('chat_id')
                if chat_id is None:
                    logger.warning("chat_id is required but not provided")
                    await websocket.send_json({"error": "chat_id is required"})
                    continue

                chat_history = chat_history_for_chain.messages(chat_id)
                formatted_history = [{"user_type": "human" if isinstance(msg, HumanMessage) else "ai" if isinstance(msg,
                                                                                                                    AIMessage) else "system",
                                      "message": msg.content} for msg in chat_history.messages]
                logger.info(f"Sending chat history: {formatted_history}")
                await websocket.send_json({"chat_history": formatted_history})

            elif action == "send_message":
                question_data = data.get('question_data')
                chat_id = data.get('chat_id')

                if question_data is None or chat_id is None:
                    logger.warning("Question data and chat_id are required but not provided")
                    await websocket.send_json({"error": "Question data and chat_id are required"})
                    continue

                question = question_data.get('question')
                if question is None:
                    logger.warning("Question is required but not provided")
                    await websocket.send_json({"error": "Question is required"})
                    continue

                logger.info(f"Processing question: {question} for chat_id: {chat_id}")

                try:
                    answer = chain_with_message_history.invoke(
                        {"question": question, "context": format_docs(retriever.invoke(question))},
                        {"configurable": {"session_id": chat_id}}
                    ).content
                    logger.info(f"Generated answer: {answer}")

                    chat_history_for_chain.add_message(HumanMessage(content=question), chat_id)
                    chat_history_for_chain.add_message(AIMessage(content=answer), chat_id)

                    logger.info(f"Sending answer: {answer}")
                    await websocket.send_json({"answer": answer})
                except Exception as e:
                    logger.error(f"Error generating answer: {str(e)}", exc_info=True)
                    await websocket.send_json({"error": str(e)})
            else:
                logger.warning(f"Unknown action: {action}")
                await websocket.send_json({"error": "Unknown action"})
    except WebSocketDisconnect:
        logger.info("Client disconnected")
    except Exception as e:
        logger.error(f"Unexpected error in WebSocket handler: {str(e)}", exc_info=True)

class ChatMetadata(BaseModel):
    chat_id: str
    chat_name: str

class SaveMessageRequest(BaseModel):
    chat_id: str
    chat_name: Optional[str]
    user_type: str
    message: str


@app.post("/save_chat")
async def save_chat(metadata: ChatMetadata):
    try:
        with sqlite3.connect('metadata.db') as conn:
            conn.execute("INSERT OR REPLACE INTO chats (chat_id, chat_name) VALUES (?, ?)",
                         (metadata.chat_id, metadata.chat_name))
            conn.commit()
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Chat metadata could not be saved")
    return {"status": "success"}
@app.post("/save_message")
async def save_message(request: SaveMessageRequest):
    chat_history_for_chain.add_message(
        HumanMessage(content=request.message) if request.user_type == 'human' else AIMessage(content=request.message),
        request.chat_id
    )
    return {"status": "success"}

@app.get("/chat_history/{chat_id}", response_class=JSONResponse)
async def get_chat_history(chat_id: str):
    history = chat_history_for_chain.messages(chat_id)
    return JSONResponse(content=[{"user_type": msg.__class__.__name__.replace("Message", "").lower(), "message": msg.content} for msg in history.messages])
@app.get("/forgot-password", response_class=HTMLResponse)
async def get_forgot_password():
    return FileResponse("static/forgot_password.html")

@app.get("/get_chat_list")
async def get_chat_list():
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT chat_id, chat_name FROM chats")
        chats = [{"id": row[0], "title": row[1]} for row in cursor.fetchall()]
    return JSONResponse(content={"chats": chats})

@app.post("/create_new_chat")
async def create_new_chat():
    chat_id = f'chat-{int(time.time() * 1000)}'
    chat_name = "Новый чат"
    with sqlite3.connect('metadata.db') as conn:
        conn.execute("INSERT INTO chats (chat_id, chat_name) VALUES (?, ?)", (chat_id, chat_name))
        conn.commit()

    # Добавляем приветственное сообщение
    welcome_message = "Вас приветствует А100! Напишите Ваш вопрос о документообороте."
    chat_history_for_chain.add_message(SystemMessage(content=welcome_message), chat_id)

    return JSONResponse(content={"chat_id": chat_id, "chat_name": chat_name})

@app.post("/forgot-password")
async def post_forgot_password(email: str = Form(...)):
    new_password = generate_random_password()
    hashed_password = get_password_hash(new_password)
    with sqlite3.connect('metadata.db') as conn:
        cursor = conn.cursor()
        cursor.execute("UPDATE users SET hashed_password = ? WHERE username = ?", (hashed_password, email))
        conn.commit()
    send_reset_password_email(email, new_password)
    return HTMLResponse("Новый пароль был отправлен на ваш email")

@app.delete("/delete_chat/{chat_id}")
async def delete_chat(chat_id: str):
    with sqlite3.connect('metadata.db') as conn:
        conn.execute("DELETE FROM chats WHERE chat_id = ?", (chat_id,))
        conn.execute("DELETE FROM history_messages WHERE chat_id = ?", (chat_id,))
        conn.commit()
    return JSONResponse(content={"status": "success"})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8222))
    uvicorn.run(app, host="0.0.0.0", port=port)
